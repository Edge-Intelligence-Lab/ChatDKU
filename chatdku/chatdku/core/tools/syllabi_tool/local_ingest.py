#!/usr/bin/env python3
"""
Local PDF/DOCX Ingestor
A command line utility that extracts structured data from PDFs and DOCX files
using SGLang hosted Qwen3 model and stores results in PostgreSQL database.
"""

import argparse
import os
import sys
import json
import logging
import getpass
from pathlib import Path
from datetime import datetime
import hashlib
from typing import Dict, Any, List, Optional, Union
import re

# Third-party imports
import psycopg2
import psycopg2.extras
from psycopg2.extensions import ISOLATION_LEVEL_AUTOCOMMIT
import pymupdf  # PyMuPDF for PDF parsing
import pdfplumber  # Alternative PDF parser for complex layouts
from docx import Document  # python-docx for DOCX parsing
import requests
from jsonschema import validate, ValidationError
import dspy

def remove_think_section(text: str) -> str:
    """
    Removes the first <think>...</think> section (including the tags) from the string.
    Works across multiple lines.
    """
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)


# Database connection configuration (hardcoded as requested)
DEFAULT_DB_CONFIG = {
    "host": "localhost",
    "port": 5432,
    "database": "chatdku_db",
    "user": "chatdku_user",
}


class DocumentIngestor:
    """Main class for document ingestion and processing"""

    def __init__(self, args):
        self.args = args
        self.setup_logging()
        self.setup_database_connection()
        self.setup_sglang_client()
        self.load_schema()

    def setup_logging(self):
        """Setup logging to file in the pool directory"""
        log_file = Path(self.args.pool) / "ingestor.log"

        # Create directory if it doesn't exist
        log_file.parent.mkdir(parents=True, exist_ok=True)

        # Setup logging
        logging.basicConfig(
            level=logging.INFO,
            format="%(asctime)s - %(levelname)s - %(message)s",
            handlers=[
                logging.FileHandler(log_file, mode="a"),
                logging.StreamHandler(sys.stdout),
            ],
        )
        self.logger = logging.getLogger(__name__)
        self.logger.info("Document Ingestor started")

    def setup_database_connection(self):
        """Setup PostgreSQL database connection with secure password handling"""
        db_config = DEFAULT_DB_CONFIG.copy()

        # Override with command line arguments if provided
        if self.args.db_host:
            db_config["host"] = self.args.db_host
        if self.args.db_port:
            db_config["port"] = self.args.db_port
        if self.args.db_name:
            db_config["database"] = self.args.db_name
        if self.args.db_user:
            db_config["user"] = self.args.db_user

        # Get password from environment variable or prompt
        password = os.environ.get("DB_PWD")
        if not password:
            password = getpass.getpass(
                f"Enter password for PostgreSQL user '{db_config['user']}': "
            )

        db_config["password"] = password

        try:
            self.conn = psycopg2.connect(**db_config)
            self.conn.set_isolation_level(ISOLATION_LEVEL_AUTOCOMMIT)
            self.logger.info(
                f"Connected to PostgreSQL database: {db_config['database']}"
            )
        except psycopg2.Error as e:
            self.logger.error(f"Failed to connect to database: {e}")
            sys.exit(1)

    def setup_sglang_client(self):
        """Setup SGLang client for Qwen3 model"""
        # SGLang serves models via OpenAI-compatible API
        new_lm = dspy.OpenAI(
            model="Qwen/Qwen3-8B",
            api_base=self.args.sglang_url,
            api_key="dummy",
            model_type="chat",
            max_tokens=40960,
            stop=["<|im_end|>"],
            temperature=0.1,
            system_prompt="You must extract structured data from documents based on provided JSON schema, adhering strictly to the schema without adding or omitting required fields.",
        )
        dspy.configure(lm=new_lm)
        self.logger.info(f"SGLang client configured for: {self.args.sglang_url}")

    def load_schema(self):
        """Load and validate JSON schema"""
        try:
            with open(self.args.schema, "r") as f:
                self.schema = json.load(f)
            self.logger.info(f"Schema loaded from: {self.args.schema}")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            self.logger.error(f"Failed to load schema: {e}")
            sys.exit(1)

    def get_file_hash(self, file_path: Path) -> str:
        """Generate MD5 hash of file for tracking processed documents"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def is_already_processed(self, file_path: Path) -> bool:
        """Check if document was already processed by reading log file"""
        log_file = Path(self.args.pool) / "ingestor.log"
        if not log_file.exists():
            return False

        file_hash = self.get_file_hash(file_path)
        try:
            with open(log_file, "r") as f:
                content = f.read()
                return f"PROCESSED: {file_path.name} (hash: {file_hash})" in content
        except Exception as e:
            self.logger.warning(f"Could not read log file: {e}")
            return False

    def mark_as_processed(self, file_path: Path):
        """Mark document as processed in log file"""
        file_hash = self.get_file_hash(file_path)
        self.logger.info(f"PROCESSED: {file_path.name} (hash: {file_hash})")

    def extract_pdf_content(self, file_path: Path) -> str:
        """Extract text content from PDF using PyMuPDF with pdfplumber fallback"""
        text_content = ""

        try:
            # Primary method: PyMuPDF (faster)
            doc = pymupdf.open(file_path)
            for page in doc:
                text_content += page.get_textpage().extractText()
            doc.close()
            self.logger.info(f"Extracted text using PyMuPDF from {file_path.name}")

        except Exception as e:
            self.logger.warning(f"PyMuPDF failed for {file_path.name}: {e}")

            try:
                # Fallback method: pdfplumber (better for complex layouts)
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += page_text + "\n"
                self.logger.info(
                    f"Extracted text using pdfplumber from {file_path.name}"
                )

            except Exception as e2:
                self.logger.error(f"Both PDF parsers failed for {file_path.name}: {e2}")
                return ""

        return text_content.strip()

    def extract_docx_content(self, file_path: Path) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = Document(str(file_path))
            text_content = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            result = "\n".join(text_content)
            self.logger.info(f"Extracted text from DOCX file: {file_path.name}")
            return result

        except Exception as e:
            self.logger.error(
                f"Failed to extract DOCX content from {file_path.name}: {e}"
            )
            return ""

    def extract_structured_data(
        self, content: str, file_name: str
    ) -> Optional[Dict[str, Any]]:
        """Use SGLang + Qwen3 to extract structured data from content"""

        # Create prompt for structured extraction based on schema
        schema_description = json.dumps(self.schema, indent=2)

        try:
            # response = dspy.Predict("json_schema, parsed_content -> extracted_json")(
            #     json_schema=schema_description, parsed_content=content
            # ).extracted_json

            class Extractor(dspy.Signature):
                """json_schema, parsed_content -> extracted_json"""
                parsed_content: str = dspy.InputField(desc="A parsed plaintext representation of a Duke Kunshan University syllabus.")
                json_schema: str = dspy.InputField(desc="A v7 json-schema description of the structured data required for syllabus information extraction.")
                extracted_json: str = dspy.OutputField(desc="A non-markdown, pure JSON reproduction of the syllabus data.")

            extractor = dspy.Predict(Extractor)
            json_text = extractor(
                parsed_content=content,
                json_schema=schema_description,
                ).extracted_json

            json_text = remove_think_section(json_text)
            # Clean up response (remove markdown formatting if present)
            if json_text.startswith("```json"):
                json_text = json_text[7:]
            if json_text.endswith("```"):
                json_text = json_text[:-3]

            # Parse and validate JSON
            self.logger.info(f"LLM response for {file_name}:\n{json_text}")
            extracted_data = json.loads(json_text.strip())

            # Validate against schema
            # validate(instance=extracted_data, schema=self.schema)

            # Add metadata
            # extracted_data["_metadata"] = {
            #     "source_file": file_name,
            #     "extraction_timestamp": datetime.now().isoformat(),
            #     "model_used": self.args.model_name,
            # }

            self.logger.info(f"Successfully extracted structured data from {file_name}")
            return extracted_data

        except (json.JSONDecodeError, ValidationError) as e:
            self.logger.error(
                f"Invalid JSON or schema validation failed for {file_name}: {e}"
            )
            return None
        except Exception as e:
            self.logger.error(f"LLM extraction failed for {file_name}: {e}")
            return None

    def store_in_database(self, data: Dict[str, Any], table_name: str = "classes"):
        """Store extracted data in PostgreSQL database, dynamically handling columns and values."""
        if not table_name:
            table_name = self.args.table_name

        try:
            from psycopg2 import sql
            self.logger.info("Creating cursor.")
            cursor = self.conn.cursor()

            # Prepare columns and values
            columns = list(data.keys())
            values = [data[col] for col in columns]

            # Build SQL statement safely
            insert_sql = sql.SQL("INSERT INTO {table} ({fields}) VALUES ({placeholders})").format(
                table=sql.Identifier(table_name),
                fields=sql.SQL(", ").join(map(sql.Identifier, columns)),
                placeholders=sql.SQL(", ").join(sql.Placeholder() * len(columns))
            )

            self.logger.info(f"Executing SQL: {insert_sql.as_string(cursor)}")
            cursor.execute(insert_sql, values)

            self.logger.info(
                f"Data stored in database table '{table_name}'"
            )
            cursor.close()

        except psycopg2.Error as e:
            self.logger.error(f"Database storage failed: {e}")

    def process_file(self, file_path: Path):
        """Process a single document file"""
        self.logger.info(f"Processing file: {file_path}")

        # Check if already processed
        if self.is_already_processed(file_path):
            self.logger.info(f"Skipping already processed file: {file_path.name}")
            return

        # Extract content based on file type
        if file_path.suffix.lower() == ".pdf":
            content = self.extract_pdf_content(file_path)
        elif file_path.suffix.lower() == ".docx":
            content = self.extract_docx_content(file_path)
        else:
            self.logger.warning(f"Unsupported file type: {file_path.suffix}")
            return

        if not content:
            self.logger.error(f"No content extracted from {file_path.name}")
            return

        # Extract structured data using SGLang + Qwen3
        structured_data = self.extract_structured_data(content, file_path.name)

        if structured_data:
            # Store in database (commit one by one as requested)
            self.store_in_database(structured_data)

            # Mark as processed
            self.mark_as_processed(file_path)
        else:
            self.logger.error(
                f"Failed to extract structured data from {file_path.name}"
            )

    def process_pool(self):
        """Process all PDF and DOCX files recursively in the pool directory and its subdirectories"""
        pool_path = Path(self.args.pool)

        if not pool_path.exists():
            self.logger.error(f"Pool directory does not exist: {pool_path}")
            sys.exit(1)

        # Find all PDF and DOCX files recursively
        pdf_files = list(pool_path.rglob("*.pdf")) + list(pool_path.rglob("*.PDF"))
        docx_files = list(pool_path.rglob("*.docx")) + list(pool_path.rglob("*.DOCX"))

        all_files = pdf_files + docx_files

        if not all_files:
            self.logger.info(
                "No PDF or DOCX files found in pool directory or its subdirectories"
            )
            return

        self.logger.info(f"Found {len(all_files)} files to process")

        # Process each file
        import traceback
        for file_path in all_files:
            try:
                self.process_file(file_path)
            except Exception as e:
                tb_str = traceback.format_exc()
                self.logger.error(f"Unexpected error processing {file_path.name}: {e}\nTraceback:\n{tb_str}")
                continue

        self.logger.info("Processing completed")

    def cleanup(self):
        """Clean up resources"""
        if hasattr(self, "conn"):
            self.conn.close()
            self.logger.info("Database connection closed")


def create_default_schema():
    """Create a default schema.json file if it doesn't exist"""
    default_schema = {
        "$schema": "http://json-schema.org/draft-07/schema#",
        "title": "Document Extraction Schema",
        "description": "Schema for structured document data extraction",
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Document title"},
            "content": {"type": "string", "description": "Main document content"},
            "metadata": {
                "type": "object",
                "properties": {
                    "author": {"type": "string"},
                    "date": {"type": "string"},
                    "category": {"type": "string"},
                },
            },
        },
        "required": ["title", "content"],
    }

    schema_file = Path("schema.json")
    if not schema_file.exists():
        with open(schema_file, "w") as f:
            json.dump(default_schema, f, indent=2)
        print(f"Created default schema file: {schema_file}")


def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(
        description="Extract structured data from PDFs and DOCX files using SGLang + Qwen3",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python local_ingestor.py
  python local_ingestor.py --pool /path/to/pdfs --schema custom_schema.json
  python local_ingestor.py --db-host remote-db --db-port 5432 --table-name documents
        """,
    )

    # File and directory arguments
    parser.add_argument(
        "--pool",
        default="./pdf_pool/",
        help="Directory containing PDF/DOCX files to process (default: ./pdf_pool/)",
    )

    parser.add_argument(
        "--schema",
        default="classes_schema.json",
        help="JSON schema file for data extraction (default: schema.json)",
    )

    # Database connection arguments
    parser.add_argument(
        "--db-host", help=f"PostgreSQL host (default: {DEFAULT_DB_CONFIG['host']})"
    )

    parser.add_argument(
        "--db-port",
        type=int,
        help=f"PostgreSQL port (default: {DEFAULT_DB_CONFIG['port']})",
    )

    parser.add_argument(
        "--db-name",
        help=f"PostgreSQL database name (default: {DEFAULT_DB_CONFIG['database']})",
    )

    parser.add_argument(
        "--db-user", help=f"PostgreSQL username (default: {DEFAULT_DB_CONFIG['user']})"
    )

    parser.add_argument(
        "--table-name",
        default="documents",
        help="Database table name for storing extracted data (default: documents)",
    )

    # SGLang configuration
    parser.add_argument(
        "--sglang-url",
        default="http://localhost:18085/v1",
        help="SGLang server URL (default: http://localhost:18085/v1)",
    )

    parser.add_argument(
        "--model-name",
        default="Qwen/Qwen3-8B",
        help="Model name for SGLang (default: Qwen/Qwen3-8B)",
    )

    # Utility arguments
    parser.add_argument(
        "--create-schema",
        action="store_true",
        help="Create a default schema.json file and exit",
    )

    parser.add_argument("--version", action="version", version="Local Ingestor v1.0.0")

    args = parser.parse_args()

    # Handle special actions
    if args.create_schema:
        create_default_schema()
        return

    # Main processing
    ingestor = None
    try:
        ingestor = DocumentIngestor(args)
        ingestor.process_pool()
    except KeyboardInterrupt:
        print("\nOperation cancelled by user")
        sys.exit(1)
    except Exception as e:
        print(f"Unexpected error: {e}")
        sys.exit(1)
    finally:
        if ingestor:
            ingestor.cleanup()


if __name__ == "__main__":
    main()
